# Django Imports
from django import forms
from django.conf import settings
from django.contrib import messages
from django.contrib.auth import authenticate, login as auth_login, logout as auth_logout, update_session_auth_hash
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.core.mail import send_mail, EmailMessage
from django.db import IntegrityError, transaction
from django.db.models import Sum, F, ExpressionWrapper, DecimalField, Count, Prefetch, Q
from django.db.models.functions import Coalesce
from django.http import JsonResponse, HttpResponseBadRequest, HttpResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.template.response import TemplateResponse
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST

# Python Standard Library Imports
import calendar
import csv
import io
import json
import locale
import mimetypes
import os
import random
import secrets
import time
import uuid
from datetime import date, timedelta, datetime
from decimal import Decimal

# Third-Party Imports
import requests
from captcha.fields import CaptchaField
from docx import Document
from geopy.exc import GeocoderTimedOut, GeocoderUnavailable
from geopy.geocoders import Nominatim
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, NamedStyle
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph

# Local Application Imports
from .decorators import add_session_timeout_context
from .models import Perfil, PasswordResetCode, Produto, Fornecedor, Cliente, Venda, HistoricoNotaFiscal

# Configura o locale para formatação de moeda
try:
    locale.setlocale(locale.LC_ALL, 'pt_BR.utf8')
except locale.Error:
    try:
        locale.setlocale(locale.LC_ALL, 'Portuguese_Brazil')
    except locale.Error:
        pass
# Formulário de Login com Captcha
class LoginFormWithCaptcha(forms.Form):
    login_usuario = forms.CharField()
    password = forms.CharField(widget=forms.PasswordInput)
    captcha = CaptchaField()
# --- VIEWS DE PÁGINAS E AUTENTICAÇÃO ---
def login_page(request):
    if request.user.is_authenticated:
        return redirect('inicio')
    login_attempts = request.session.get('login_attempts', 0)
    show_captcha = login_attempts >= 3
    
    if request.method == 'POST':
        form = LoginFormWithCaptcha(request.POST) if show_captcha else None
        
        if show_captcha and form and not form.is_valid():
            request.session['login_attempts'] = login_attempts + 1
            messages.error(request, 'O texto do CAPTCHA está incorreto.')
            return redirect('login')
        
        username = request.POST.get('login_usuario')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        if user is not None:
            # SUCESSO: Zera o contador, faz o login e redireciona.
            if 'login_attempts' in request.session:
                del request.session['login_attempts']
            auth_login(request, user)
            return redirect('inicio')
        else:
            # FALHA: Incrementa o contador.
            request.session['login_attempts'] = login_attempts + 1
            messages.error(request, 'Nome de usuário ou senha incorretos.')
            return redirect('login')
    
    # Para requisições GET (novas ou redirecionadas)
    form = LoginFormWithCaptcha() if show_captcha else None
    return render(request, 'login.html', {
        'show_captcha': show_captcha,
        'captcha_form': form
    })
def user_logout(request):
    auth_logout(request)
    # Garante que o contador seja limpo ao sair do sistema
    if 'login_attempts' in request.session:
        del request.session['login_attempts']
    return redirect('login')
def senha(request):
    return render(request, 'senha.html')
def cadastrar(request):
    # Esta view pode permanecer como está pois é pública
    if request.method == 'POST':
        # ... seu código de cadastro ...
        nomecompleto = request.POST.get('nomecompleto', '').strip()
        username = request.POST.get('usuario', '').strip()
        email = request.POST.get('email', '').strip()
        senha = request.POST.get('senha', '').strip()
        repetir_senha = request.POST.get('repetir_senha', '').strip()
        celular = request.POST.get('celular', '').strip()
        if not all([nomecompleto, username, email, senha, repetir_senha, celular]):
            return JsonResponse({'success': False, 'message': 'Por favor, preencha todos os campos obrigatórios.'})
        if senha != repetir_senha:
            return JsonResponse({'success': False, 'message': 'As senhas não coincidem. Tente novamente.'})
        if User.objects.filter(username=username).exists():
            return JsonResponse({'success': False, 'message': 'Este nome de usuário já está em uso.'})
        if User.objects.filter(email=email).exists():
            return JsonResponse({'success': False, 'message': 'Este e-mail já está em uso.'})
        if Perfil.objects.filter(celular=celular).exists():
            return JsonResponse({'success': False, 'message': 'Este número de celular já está em uso.'})
        try:
            user = User.objects.create_user(username=username, password=senha, email=email)
            user.first_name = nomecompleto.split(' ')[0]
            user.last_name = ' '.join(nomecompleto.split(' ')[1:])
            user.save()
            Perfil.objects.create(user=user, celular=celular)
            return JsonResponse({'success': True, 'message': 'Usuário cadastrado com sucesso!'})
        except Exception as e:
            return JsonResponse({'success': False, 'message': f'Erro ao criar usuário: {str(e)}'})
    return render(request, 'cadastrar.html')
@csrf_exempt
@require_POST
def recuperar_senha(request):
    # ... seu código de recuperação ...
    try:
        data = json.loads(request.body.decode('utf-8'))
        email = data.get('email', '').strip()
        if not email:
            return JsonResponse({'success': False, 'message': 'E-mail é obrigatório.'})
        user = None
        try:
            user = User.objects.get(email=email)
        except User.DoesNotExist:
            return JsonResponse({'success': False, 'message': 'E-mail não encontrado.'})
        codigo = ''.join(random.choices('0123456789', k=6))
        PasswordResetCode.objects.filter(user=user).delete()
        PasswordResetCode.objects.create(user=user, code=codigo)
        send_mail(
            'Código de Recuperação de Senha',
            f'Olá, seu código de recuperação é: {codigo}\nEste código é válido por 10 minutos.',
            settings.EMAIL_HOST_USER,
            [user.email],
            fail_silently=False,
        )
        return JsonResponse({
            'success': True,
            'message': 'Um código de recuperação foi enviado para o seu e-mail.'
        })
    except json.JSONDecodeError:
        return HttpResponseBadRequest('JSON inválido.')
    except Exception as e:
        print(f"Erro ao enviar e-mail: {e}")
        return JsonResponse({
            'success': False,
            'message': 'Ocorreu um erro ao processar sua solicitação. Tente novamente mais tarde.'
        })
@csrf_exempt
@require_POST
def validar_codigo(request):
    # ... seu código de validação ...
    try:
        data = json.loads(request.body.decode('utf-8'))
        email = data.get('email', '').strip()
        codigo = data.get('codigo', '').strip()
        if not email or not codigo:
            return JsonResponse({'success': False, 'message': 'E-mail e código são obrigatórios.'})
        try:
            user = User.objects.get(email=email)
        except User.DoesNotExist:
            return JsonResponse({'success': False, 'message': 'E-mail não encontrado.'})
        reset_code_obj = PasswordResetCode.objects.filter(user=user, code=codigo).order_by('-created_at').first()
        if reset_code_obj and reset_code_obj.is_valid():
            reset_code_obj.delete()
            return JsonResponse({'success': True, 'message': 'Código validado com sucesso.'})
        else:
            return JsonResponse({'success': False, 'message': 'Código inválido ou expirado.'})
    except json.JSONDecodeError:
        return HttpResponseBadRequest('JSON inválido.')
    except Exception as e:
        print(f"Erro ao validar código: {e}")
        return JsonResponse({
            'success': False,
            'message': 'Ocorreu um erro ao processar sua solicitação. Tente novamente.'
        })
@csrf_exempt
@require_POST
def redefinir_senha(request):
    # ... seu código de redefinição ...
    try:
        data = json.loads(request.body.decode('utf-8'))
        email = data.get('email', '').strip()
        nova_senha = data.get('nova_senha', '').strip()
        if not email or not nova_senha:
            return JsonResponse({'success': False, 'message': 'E-mail e nova senha são obrigatórios.'})
        try:
            user = User.objects.get(email=email)
        except User.DoesNotExist:
            return JsonResponse({'success': False, 'message': 'E-mail não encontrado.'})
        user.set_password(nova_senha)
        user.save()
        return JsonResponse({'success': True, 'message': 'Senha alterada com sucesso. Você pode fazer login agora.'})
    except json.JSONDecodeError:
        return HttpResponseBadRequest('JSON inválido.')
    except Exception as e:
        print(f"Erro ao redefinir a senha: {e}")
        return JsonResponse({
            'success': False,
            'message': 'Ocorreu um erro ao redefinir a senha. Tente novamente.'
        })
# --- VIEWS DE PÁGINAS PROTEGIDAS ---
@login_required
@add_session_timeout_context
def inicio(request):
    return TemplateResponse(request, 'inicio.html', {})
@login_required
@add_session_timeout_context
def clientes(request):
    return TemplateResponse(request, 'clientes.html', {})
@login_required
@add_session_timeout_context
def fornecedores(request):
    return TemplateResponse(request, 'fornecedores.html', {})
@login_required
@add_session_timeout_context
def vendas(request):
    return TemplateResponse(request, 'vendas.html', {})
@login_required
@add_session_timeout_context
def relatorios(request):
    return TemplateResponse(request, 'relatorios.html', {})
@login_required
@add_session_timeout_context
def produtos(request):
    return TemplateResponse(request, 'produtos.html', {})
@login_required
@add_session_timeout_context
def configuracoes(request):
    return TemplateResponse(request, 'configuracoes.html', {})
@login_required
def barra(request):
    return render(request, 'barra.html')
# --- VIEWS DE API PROTEGIDAS ---
@require_POST
@login_required
def api_reauthenticate(request):
    # ... seu código de reautenticação ...
    try:
        data = json.loads(request.body)
        password = data.get('password')
        if not password:
            return JsonResponse({'success': False, 'error': 'Senha não fornecida.'}, status=400)
    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'error': 'Requisição inválida.'}, status=400)
    user = request.user
    if user.check_password(password):
        return JsonResponse({'success': True})
    else:
        return JsonResponse({'success': False, 'error': 'Senha incorreta.'}, status=401)
    
@login_required
def api_user_profile(request):
    # ... seu código de perfil ...
    if request.method == 'GET':
        user = request.user
        try:
            perfil = Perfil.objects.get(user=user)
        except Perfil.DoesNotExist:
            perfil = Perfil.objects.create(user=user)
        user_data = {
            'username': user.username,
            'fullname': f"{user.first_name} {user.last_name}".strip(),
            'email': user.email,
            'celular': perfil.celular,
        }
        return JsonResponse(user_data)
    return JsonResponse({'error': 'Método não permitido.'}, status=405)
@require_POST
@login_required
def api_enviar_codigo_alteracao_senha(request):
    # ... seu código de enviar código ...
    try:
        data = json.loads(request.body.decode('utf-8'))
        current_password = data.get('current_password')
        user = request.user
        email = user.email
        if not user.check_password(current_password):
            return JsonResponse({'success': False, 'message': 'Senha atual incorreta.'}, status=400)
        if not email:
            return JsonResponse({'success': False, 'message': 'E-mail não cadastrado para o usuário.'}, status=400)
        codigo = ''.join(random.choices('0123456789', k=6))
        PasswordResetCode.objects.update_or_create(
            user=user,
            defaults={'code': codigo, 'created_at': timezone.now()}
        )
        send_mail(
            'Código de Alteração de Senha',
            f'Olá, seu código para alterar a senha é: {codigo}\nEste código é válido por 10 minutos.',
            settings.EMAIL_HOST_USER,
            [email],
            fail_silently=False,
        )
        return JsonResponse({'success': True, 'message': 'Um código de validação foi enviado para o seu e-mail.'})
    except Exception as e:
        print(f"Erro ao enviar código de alteração de senha: {e}")
        return JsonResponse({'success': False, 'message': 'Ocorreu um erro. Tente novamente mais tarde.'}, status=500)
@require_POST
# REMOVA a sua função "api_user_profile" antiga.
# SUBSTITUA a sua função "api_update_user_profile" por esta versão unificada abaixo.
# Ela agora lida com GET (buscar perfil) e POST (atualizar perfil).

@login_required
def api_update_user_profile(request):
    # Lógica para buscar os dados do perfil (GET)
    if request.method == 'GET':
        user = request.user
        try:
            perfil = Perfil.objects.get(user=user)
        except Perfil.DoesNotExist:
            # Cria um perfil se não existir, para garantir consistência
            perfil = Perfil.objects.create(user=user)
        
        user_data = {
            'username': user.username,
            'fullname': f"{user.first_name} {user.last_name}".strip(),
            'email': user.email,
            'celular': perfil.celular,
        }
        return JsonResponse(user_data)

    # Lógica para atualizar os dados do perfil (POST)
    if request.method == 'POST':
        try:
            data = json.loads(request.body.decode('utf-8'))
            user = request.user
            
            # Bloco para ALTERAR SENHA
            if 'current_password' in data and 'new_password' in data and 'code' in data:
                current_password = data.get('current_password')
                new_password = data.get('new_password')
                code = data.get('code')
                
                if not user.check_password(current_password):
                    return JsonResponse({'success': False, 'message': 'Senha atual incorreta.'}, status=400)
                
                reset_code_obj = PasswordResetCode.objects.filter(user=user, code=code).order_by('-created_at').first()
                if not (reset_code_obj and reset_code_obj.is_valid()):
                    return JsonResponse({'success': False, 'message': 'Código de validação inválido ou expirado.'}, status=400)

                if len(new_password) < 8:
                    return JsonResponse({'success': False, 'message': 'A nova senha deve ter pelo menos 8 caracteres.'}, status=400)
                
                with transaction.atomic():
                    user.set_password(new_password)
                    user.save()
                    reset_code_obj.delete()
                
                update_session_auth_hash(request, user)
                return JsonResponse({'success': True, 'message': 'Senha alterada com sucesso!'})

            # Bloco para ATUALIZAR INFORMAÇÕES PESSOAIS
            elif 'fullname' in data or 'email' in data or 'celular' in data:
                fullname = data.get('fullname')
                email = data.get('email')
                celular = data.get('celular')
                
                with transaction.atomic():
                    try:
                        perfil, created = Perfil.objects.get_or_create(user=user)
                        if fullname:
                            partes_nome = fullname.strip().split(' ', 1)
                            user.first_name = partes_nome[0]
                            user.last_name = partes_nome[1] if len(partes_nome) > 1 else ''
                        
                        if email and email != user.email:
                            if User.objects.filter(email=email).exclude(pk=user.pk).exists():
                                return JsonResponse({'success': False, 'message': 'Este e-mail já está em uso.'}, status=400)
                            user.email = email
                        
                        if celular and celular != perfil.celular:
                            if Perfil.objects.filter(celular=celular).exclude(pk=perfil.pk).exists():
                                return JsonResponse({'success': False, 'message': 'Este número de celular já está em uso.'}, status=400)
                            perfil.celular = celular
                        
                        user.save()
                        perfil.save()
                        return JsonResponse({'success': True, 'message': 'Informações pessoais atualizadas com sucesso!'})
                    
                    except Exception as e:
                        return JsonResponse({'success': False, 'message': f'Erro ao atualizar informações: {str(e)}'}, status=400)
            
            else:
                return JsonResponse({'success': False, 'message': 'Dados de requisição inválidos.'}, status=400)

        except json.JSONDecodeError:
            return HttpResponseBadRequest('JSON inválido.')
        except Exception as e:
            return JsonResponse({'success': False, 'message': f'Erro inesperado: {str(e)}'}, status=500)

    # Se o método não for GET nem POST
    return JsonResponse({'error': 'Método não permitido.'}, status=405)
@login_required
def api_produtos(request):
    # ... seu código de produtos ...
    if request.method == 'GET':
        search_query = request.GET.get('search', None)
        produtos_queryset = Produto.objects.all().select_related('fornecedor').prefetch_related(
            Prefetch('historico_notas', queryset=HistoricoNotaFiscal.objects.order_by('-data_entrada'), to_attr='last_note_list')
        )
        if search_query:
            produtos_queryset = produtos_queryset.filter(
                Q(nome__icontains=search_query) |
                Q(id__icontains=search_query) |
                Q(fornecedor__nome__icontains=search_query)
            )
        produtos_list = []
        for produto in produtos_queryset:
            fornecedor_data = {}
            if produto.fornecedor:
                fornecedor_data = {'nome': produto.fornecedor.nome, 'cnpj': produto.fornecedor.cnpj}
            last_note_data = None
            if hasattr(produto, 'last_note_list') and produto.last_note_list:
                last_note_data = {'numero_nota': produto.last_note_list[0].numero_nota}
            produtos_list.append({
                'id': produto.id, 'nome': produto.nome, 'quantidade': produto.quantidade,
                'quantidade_minima': produto.quantidade_minima, 'preco_compra': produto.preco_compra,
                'preco_venda': produto.preco_venda, 'fornecedor': fornecedor_data,
                'last_nota_fiscal': last_note_data,
            })
        return JsonResponse(produtos_list, safe=False)
    elif request.method == 'POST':
        try:
            data = json.loads(request.body.decode('utf-8'))
            fornecedor = get_object_or_404(Fornecedor, cnpj=data.get('fornecedor_cnpj'))
            with transaction.atomic():
                produto = Produto.objects.create(
                    nome=data.get('nome'),
                    quantidade=data.get('quantidade'),
                    quantidade_minima=data.get('quantidade_minima', 0),
                    preco_compra=data.get('preco_compra'),
                    preco_venda=data.get('preco_venda'),
                    fornecedor=fornecedor
                )
                HistoricoNotaFiscal.objects.create(
                    produto=produto,
                    numero_nota=data.get('nota_fiscal'),
                    quantidade_adicionada=data.get('quantidade'),
                    preco_compra_nota=data.get('preco_compra'),
                    preco_venda_nota=data.get('preco_venda')
                )
            return JsonResponse({'message': 'Produto criado com sucesso!'}, status=201)
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)
@login_required
def api_produto_detalhe(request, id):
    # ... seu código de detalhe do produto ...
    try:
        produto = Produto.objects.get(id=id)
    except Produto.DoesNotExist:
        return JsonResponse({'error': 'Produto não encontrado.'}, status=404)
    if request.method == 'PUT':
        try:
            data = json.loads(request.body.decode('utf-8'))
            fornecedor_cnpj = data.get('fornecedor_cnpj')
            fornecedor = produto.fornecedor
            if fornecedor_cnpj and (fornecedor is None or fornecedor.cnpj != fornecedor_cnpj):
                fornecedor = get_object_or_404(Fornecedor, cnpj=fornecedor_cnpj)
            with transaction.atomic():
                produto.nome = data.get('nome', produto.nome)
                produto.quantidade_minima = data.get('quantidade_minima', produto.quantidade_minima)
                produto.fornecedor = fornecedor
                produto.preco_compra = data.get('preco_compra', produto.preco_compra)
                produto.preco_venda = data.get('preco_venda', produto.preco_venda)
                produto.save()
            produto_data = { 'id': produto.id, 'nome': produto.nome, 'quantidade': produto.quantidade }
            return JsonResponse({'message': 'Produto atualizado com sucesso!', 'produto': produto_data})
        except Exception as e:
            return JsonResponse({'error': f'Erro ao atualizar produto: {str(e)}'}, status=400)
    elif request.method == 'DELETE':
        produto.delete()
        return HttpResponse(status=204)
    return JsonResponse({'error': 'Método não permitido.'}, status=405)

@login_required
def api_fornecedores(request):
    if request.method == 'GET':
        fornecedores = list(Fornecedor.objects.values())
        return JsonResponse(fornecedores, safe=False)
    
    elif request.method == 'POST':
        try:
            data = json.loads(request.body.decode('utf-8'))
            
            # 1. Validação de campos obrigatórios
            required_fields = ['nome', 'cnpj', 'email', 'telefone', 'endereco', 'cep', 'cidade', 'estado']
            if not all(data.get(field) and str(data.get(field)).strip() for field in required_fields):
                return JsonResponse({'error': 'Todos os campos, incluindo o endereço completo, são obrigatórios.'}, status=400)

            # 2. Validação de campos únicos antes de salvar
            cnpj = data.get('cnpj')
            email = data.get('email')
            if Fornecedor.objects.filter(cnpj=cnpj).exists():
                return JsonResponse({'error': f'O CNPJ {cnpj} já está cadastrado.'}, status=400)
            if Fornecedor.objects.filter(email=email).exists():
                return JsonResponse({'error': f'O e-mail {email} já está cadastrado para outro fornecedor.'}, status=400)

            # 3. Criação do objeto
            fornecedor = Fornecedor(
                nome=data.get('nome'),
                cnpj=cnpj,
                email=email,
                telefone=data.get('telefone'),
                endereco=data.get('endereco'),
                cep=data.get('cep'),
                cidade=data.get('cidade'),
                estado=data.get('estado')
            )
            fornecedor.save()
            
            novo_fornecedor_data = {
                'nome': fornecedor.nome, 'cnpj': fornecedor.cnpj, 'email': fornecedor.email,
                'telefone': fornecedor.telefone, 'endereco': fornecedor.endereco, 'cep': fornecedor.cep,
                'cidade': fornecedor.cidade, 'estado': fornecedor.estado
            }
            return JsonResponse(novo_fornecedor_data, status=201)

        except IntegrityError as e:
            # Captura erros de violação de regras do banco de dados (ex: outro campo UNIQUE)
            return JsonResponse({'error': f'Erro de integridade do banco de dados: {str(e)}'}, status=400)
        except Exception as e:
            # Captura todos os outros erros
            return JsonResponse({'error': f'Ocorreu um erro inesperado: {str(e)}'}, status=400)
@login_required
def api_fornecedor_detalhe(request, cnpj):
    try:
        fornecedor = Fornecedor.objects.get(cnpj=cnpj)
    except Fornecedor.DoesNotExist:
        return JsonResponse({'error': 'Fornecedor não encontrado.'}, status=404)
    
    if request.method == 'PUT':
        try:
            data = json.loads(request.body.decode('utf-8'))
            fornecedor.nome = data.get('nome', fornecedor.nome)
            fornecedor.email = data.get('email', fornecedor.email)
            fornecedor.telefone = data.get('telefone', fornecedor.telefone)
            fornecedor.endereco = data.get('endereco', fornecedor.endereco)
            fornecedor.cep = data.get('cep', fornecedor.cep)
            fornecedor.cidade = data.get('cidade', fornecedor.cidade)
            fornecedor.estado = data.get('estado', fornecedor.estado)
            fornecedor.save()
            return JsonResponse({'message': 'Fornecedor atualizado com sucesso!', 'fornecedor': data})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)
            
    elif request.method == 'DELETE':
        fornecedor.delete()
        return HttpResponse(status=204)

@login_required
def api_clientes(request):
    if request.method == 'GET':
        clientes = list(Cliente.objects.values())
        return JsonResponse(clientes, safe=False)
    
    elif request.method == 'POST':
        try:
            data = json.loads(request.body.decode('utf-8'))
            
            # Validação de campos obrigatórios
            required_fields = ['nome', 'cpf', 'email', 'telefone', 'endereco', 'cep', 'cidade', 'estado']
            if not all(data.get(field) and str(data.get(field)).strip() for field in required_fields):
                return JsonResponse({'error': 'Todos os campos são obrigatórios.'}, status=400)

            # --- VALIDAÇÃO DE DADOS ÚNICOS (CPF, E-mail e Telefone) ---
            cpf = data.get('cpf')
            email = data.get('email')
            telefone = data.get('telefone')

            if Cliente.objects.filter(cpf=cpf).exists():
                return JsonResponse({'error': f'O CPF informado já está cadastrado.'}, status=400)
            
            if Cliente.objects.filter(email=email).exists():
                return JsonResponse({'error': f'O e-mail informado já está cadastrado.'}, status=400)

            if Cliente.objects.filter(telefone=telefone).exists():
                return JsonResponse({'error': f'O telefone informado já está cadastrado.'}, status=400)
            # --- FIM DA VALIDAÇÃO ---

            # Cria o cliente se todas as validações passarem
            cliente = Cliente.objects.create(
                nome=data.get('nome'), 
                cpf=cpf, 
                email=email,
                telefone=telefone, 
                endereco=data.get('endereco'), 
                cep=data.get('cep'),
                cidade=data.get('cidade'),
                estado=data.get('estado')
            )
            return JsonResponse({'message': 'Cliente cadastrado com sucesso!', 'cliente': data}, status=201)
        
        except Exception as e:
            return JsonResponse({'error': f'Erro ao adicionar cliente: {str(e)}'}, status=400)

@login_required
def api_cliente_detalhe(request, cpf):
    try:
        cliente = Cliente.objects.get(cpf=cpf)
    except Cliente.DoesNotExist:
        return JsonResponse({'error': 'Cliente não encontrado.'}, status=404)
    
    if request.method == 'PUT':
        try:
            data = json.loads(request.body.decode('utf-8'))
            cliente.nome = data.get('nome', cliente.nome)
            cliente.email = data.get('email', cliente.email)
            cliente.telefone = data.get('telefone', cliente.telefone)
            cliente.endereco = data.get('endereco', cliente.endereco)
            cliente.cep = data.get('cep', cliente.cep)
            cliente.cidade = data.get('cidade', cliente.cidade)
            cliente.estado = data.get('estado', cliente.estado)
            cliente.save()
            return JsonResponse({'message': 'Cliente atualizado com sucesso!', 'cliente': data})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)
            
    elif request.method == 'DELETE':
        cliente.delete()
        return HttpResponse(status=204)
        
    return JsonResponse({'error': 'Método não permitido.'}, status=405)

@login_required
def api_vendas(request, id=None):
    # ... seu código de vendas ...
    if request.method == 'GET':
        vendas = list(Venda.objects.select_related('produto', 'cliente').all().order_by('-data_venda').values(
            'id', 'quantidade', 'preco_total', 'data_venda',
            'produto__nome', 'produto__id',
            'cliente__nome', 'cliente__cpf'
        ))
        return JsonResponse(vendas, safe=False)
    elif request.method == 'POST':
        try:
            data = json.loads(request.body.decode('utf-8'))
            produto = get_object_or_404(Produto, id=data.get('produto_id'))
            cliente = get_object_or_404(Cliente, cpf=data.get('cliente_cpf'))
            quantidade_vendida = data.get('quantidade')
            with transaction.atomic():
                if produto.quantidade < quantidade_vendida:
                    return JsonResponse({'error': f'Estoque insuficiente. Disponível: {produto.quantidade}'}, status=400)
                preco_total = produto.preco_venda * quantidade_vendida
                venda = Venda.objects.create(
                    produto=produto, cliente=cliente,
                    quantidade=quantidade_vendida, preco_total=preco_total
                )
                produto.quantidade -= quantidade_vendida
                produto.save()
            return JsonResponse({'message': 'Venda registrada com sucesso!', 'venda_id': venda.id}, status=201)
        except Exception as e:
            return JsonResponse({'error': f'Erro ao registrar venda: {str(e)}'}, status=400)
    elif request.method == 'DELETE':
        venda = get_object_or_404(Venda, id=id)
        with transaction.atomic():
            produto = venda.produto
            produto.quantidade += venda.quantidade
            produto.save()
            venda.delete()
        return HttpResponse(status=204)
    return JsonResponse({'error': 'Método não permitido'}, status=405)
@login_required
def api_exportar_produtos(request):
    # 1. Cria um Workbook (o arquivo Excel em memória)
    wb = Workbook()
    ws = wb.active
    ws.title = "Relatório de Estoque"

    # 2. Define os cabeçalhos da planilha
    headers = [
        "Código", "Nome do Produto", "Quantidade", 
        "Preço de Custo (R$)", "Preço de Venda (R$)", "Fornecedor",
        "Valor Total (Custo)"
    ]
    ws.append(headers)

    # 3. Estiliza o cabeçalho
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="26A0B8", end_color="26A0B8", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center")
    
    for col_num, header_title in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment

    # 4. Busca os dados do banco de dados
    # Usamos select_related para otimizar a busca pelo fornecedor
    produtos = Produto.objects.select_related('fornecedor').all()

    # Define os estilos para as células de dados
    currency_style = NamedStyle(name='currency', number_format='R$ #,##0.00')
    wb.add_named_style(currency_style)

    # 5. Adiciona os dados dos produtos na planilha
    for produto in produtos:
        valor_total_custo = produto.quantidade * produto.preco_compra
        
        row_data = [
            str(produto.id)[:6],
            produto.nome,
            produto.quantidade,
            produto.preco_compra,
            produto.preco_venda,
            produto.fornecedor.nome if produto.fornecedor else "Sem Fornecedor",
            valor_total_custo
        ]
        ws.append(row_data)

    # 6. Ajusta a largura das colunas e aplica formatação
    for col_num in range(1, ws.max_column + 1):
        column_letter = get_column_letter(col_num)
        max_length = 0
        
        # Ajusta a largura
        for cell in ws[column_letter]:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = (max_length + 2)
        ws.column_dimensions[column_letter].width = adjusted_width

        # Aplica formato de moeda para colunas de preço
        header_name = ws.cell(row=1, column=col_num).value
        if "R$" in header_name or "Valor Total" in header_name:
            for cell in ws[column_letter][1:]:
                cell.style = currency_style

    # 7. (BÔNUS) Adiciona uma linha de resumo no final
    total_stock_value = produtos.aggregate(total=Sum(F('quantidade') * F('preco_compra')))['total'] or 0
    
    ws.append([])
    summary_row = ws.max_row + 1
    ws.cell(row=summary_row, column=5, value="Valor Total do Estoque (Custo):").font = Font(bold=True)
    ws.cell(row=summary_row, column=6, value=total_stock_value).style = currency_style
    ws.cell(row=summary_row, column=6).font = Font(bold=True)

    # 8. Prepara a resposta HTTP para fazer o download do arquivo
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )
    # Define um nome de arquivo dinâmico com a data atual
    filename = f"relatorio_estoque_{datetime.now().strftime('%Y-%m-%d')}.xlsx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    # Salva o workbook na resposta
    wb.save(response)

    return response

@login_required
def api_relatorios(request):
    # ... seu código de relatórios ...
    try:
        total_produtos_qnt = Produto.objects.aggregate(total_qnt=Coalesce(Sum('quantidade'), 0))['total_qnt']
        valor_total_estoque = Produto.objects.aggregate(
            total_valor=Coalesce(Sum(F('quantidade') * F('preco_venda')), Decimal(0), output_field=DecimalField())
        )['total_valor']
        produtos_falta = list(Produto.objects.filter(quantidade__lt=F('quantidade_minima')).values_list('nome', flat=True))
        vendas_por_produto = Venda.objects.values('produto__nome').annotate(
            total_vendido=Coalesce(Sum('quantidade'), 0)
        )
        produtos_mais_saem = [
            {"nome": item['produto__nome'], "vendas": item['total_vendido']}
            for item in vendas_por_produto.order_by('-total_vendido')[:3]
        ]
        produtos_menos_saem = [
            {"nome": item['produto__nome'], "vendas": item['total_vendido']}
            for item in vendas_por_produto.order_by('total_vendido')[:3]
        ]
        principal_fornecedor = Fornecedor.objects.annotate(
            valor_gasto=Coalesce(Sum(F('produto__historico_notas__quantidade_adicionada') * F('produto__preco_compra')), Decimal(0), output_field=DecimalField())
        ).order_by('-valor_gasto').first()
        principal_fornecedor_data = {
            "nome": principal_fornecedor.nome if principal_fornecedor else "N/A",
            "valorGasto": f"R$ {principal_fornecedor.valor_gasto:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.') if principal_fornecedor else "R$ 0,00"
        }
        top_clientes = list(Cliente.objects.annotate(
            valor_total_vendas=Coalesce(Sum('venda__preco_total'), Decimal(0))
        ).order_by('-valor_total_vendas').values_list('nome', flat=True)[:3])
        fornecedores_cadastrados = Fornecedor.objects.count()
        clientes_cadastrados = Cliente.objects.count()
        total_vendas = Venda.objects.aggregate(total_vendas=Coalesce(Sum('quantidade'), 0))['total_vendas']
        valor_ganho = Venda.objects.aggregate(total_ganho=Coalesce(Sum('preco_total'), Decimal(0), output_field=DecimalField()))['total_ganho']
        dados_relatorio = {
            'produtosEstoque': total_produtos_qnt,
            'valorEstoque': f"R$ {valor_total_estoque:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),
            'produtosFalta': produtos_falta,
            'produtosMaisSaem': produtos_mais_saem,
            'produtosMenosSaem': produtos_menos_saem,
            'principalFornecedor': principal_fornecedor_data,
            'topClientes': top_clientes,
            'fornecedores': fornecedores_cadastrados,
            'clientes': clientes_cadastrados,
            'totalVendas': total_vendas,
            'valorGanho': f"R$ {valor_ganho:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
        }
        return JsonResponse(dados_relatorio)
    except Exception as e:
        print(f"Erro na API de relatórios: {e}")
        return JsonResponse({'error': 'Erro ao gerar o relatório. Tente novamente mais tarde.'}, status=500)
# As funções abaixo são auxiliares e não são views, não precisam de decorators.
def generate_pdf(buffer, data):
    # ... seu código de gerar PDF ...
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    story = []
    styles = getSampleStyleSheet()
    story.append(Paragraph(f'<b>Relatório Completo do Sistema</b>', styles['Title']))
    story.append(Paragraph(f'<i>Gerado em: {date.today().strftime("%d/%m/%Y")}</i>', styles['Normal']))
    story.append(Paragraph('<br/><b>Visão Geral do Estoque e Vendas</b>', styles['h2']))
    overview_data = [['Métrica', 'Valor'],
                     ['Produtos em Estoque', data['produtosEstoque']],
                     ['Valor Total do Estoque', data['valorEstoque']],
                     ['Total de Vendas', data['totalVendas']],
                     ['Valor Total Ganho', data['valorGanho']]]
    story.append(Table(overview_data, colWidths=[200, 300], style=TableStyle([('BOX', (0,0), (-1,-1), 1, colors.black), ('GRID', (0,0), (-1,-1), 1, colors.black)])))
    story.append(Paragraph('<br/><b>Produtos com Estoque Crítico</b>', styles['h2']))
    if data['produtosFalta']:
        falta_data = [['Nome do Produto']] + [[item] for item in data['produtosFalta']]
        story.append(Table(falta_data, colWidths=[500], style=TableStyle([('BOX', (0,0), (-1,-1), 1, colors.black), ('GRID', (0,0), (-1,-1), 1, colors.black)])))
    else:
        story.append(Paragraph('Nenhum produto em falta. O estoque está saudável!', styles['Normal']))
    story.append(Paragraph('<br/><b>Ranking de Produtos</b>', styles['h2']))
    ranking_data = [['Posição', 'Nome do Produto', 'Vendas']]
    for i, item in enumerate(data['produtosMaisSaem']):
        ranking_data.append([f'Top {i+1}', item['nome'], item['vendas']])
    for i, item in enumerate(data['produtosMenosSaem']):
        ranking_data.append([f'Menos V. {i+1}', item['nome'], item['vendas']])
    story.append(Table(ranking_data, colWidths=[100, 200, 200], style=TableStyle([('BOX', (0,0), (-1,-1), 1, colors.black), ('GRID', (0,0), (-1,-1), 1, colors.black)])))
    story.append(Paragraph('<br/><b>Análise de Clientes e Fornecedores</b>', styles['h2']))
    clientes_data = [['Nome do Cliente']] + [[item] for item in data['topClientes']]
    story.append(Table(clientes_data, colWidths=[500], style=TableStyle([('BOX', (0,0), (-1,-1), 1, colors.black), ('GRID', (0,0), (-1,-1), 1, colors.black)])))
    fornecedor_data = [['Principal Fornecedor', 'Valor Gasto'],
                       [data['principalFornecedor']['nome'], data['principalFornecedor']['valorGasto']]]
    story.append(Table(fornecedor_data, colWidths=[250, 250], style=TableStyle([('BOX', (0,0), (-1,-1), 1, colors.black), ('GRID', (0,0), (-1,-1), 1, colors.black)])))
    doc.build(story)
    buffer.seek(0)
def generate_excel(buffer, data):
    # ... seu código de gerar Excel ...
    wb = Workbook()
    sheet1 = wb.active
    sheet1.title = "Visão Geral"
    sheet1.append(['Métrica', 'Valor'])
    sheet1.append(['Produtos em Estoque', data['produtosEstoque']])
    sheet1.append(['Valor Total do Estoque', data['valorEstoque']])
    sheet1.append(['Total de Vendas', data['totalVendas']])
    sheet1.append(['Valor Total Ganho', data['valorGanho']])
    sheet2 = wb.create_sheet(title="Produtos em Falta")
    sheet2.append(['Nome do Produto'])
    if data['produtosFalta']:
        for item in data['produtosFalta']:
            sheet2.append([item])
    sheet3 = wb.create_sheet(title="Ranking de Produtos")
    sheet3.append(['Posição', 'Nome do Produto', 'Vendas'])
    for i, item in enumerate(data['produtosMaisSaem']):
        sheet3.append([f'Top {i+1}', item['nome'], item['vendas']])
    for i, item in enumerate(data['produtosMenosSaem']):
        sheet3.append([f'Menos V. {i+1}', item['nome'], item['vendas']])
    sheet4 = wb.create_sheet(title="Clientes e Fornecedores")
    sheet4.append(['Tipo', 'Nome', 'Valor'])
    sheet4.append(['Principal Fornecedor', data['principalFornecedor']['nome'], data['principalFornecedor']['valorGasto']])
    for cliente in data['topClientes']:
        sheet4.append(['Principal Cliente', cliente, ''])
    wb.save(buffer)
    buffer.seek(0)
def generate_word(buffer, data):
    # ... seu código de gerar Word ...
    document = Document()
    document.add_heading('Relatório Completo do Sistema', 0)
    document.add_paragraph(f'Gerado em: {date.today().strftime("%d/%m/%Y")}')
    document.add_heading('Visão Geral do Estoque e Vendas', level=1)
    for key, value in [('Produtos em Estoque', data['produtosEstoque']),
                         ('Valor Total do Estoque', data['valorEstoque']),
                         ('Total de Vendas', data['totalVendas']),
                         ('Valor Total Ganho', data['valorGanho'])] :
        document.add_paragraph(f'• {key}: {value}', style='List Bullet')
    document.add_heading('Produtos com Estoque Crítico', level=1)
    if data['produtosFalta']:
        for item in data['produtosFalta']:
            document.add_paragraph(f'• {item}', style='List Bullet')
    else:
        document.add_paragraph('Nenhum produto em falta. O estoque está saudável!')
    document.add_heading('Ranking de Produtos', level=1)
    document.add_heading('Produtos que Mais Saem', level=2)
    if data['produtosMaisSaem']:
        for item in data['produtosMaisSaem']:
            document.add_paragraph(f'• {item["nome"]}: {item["vendas"]} unidades', style='List Bullet')
    document.add_heading('Produtos que Menos Saem', level=2)
    if data['produtosMenosSaem']:
        for item in data['produtosMenosSaem']:
            document.add_paragraph(f'• {item["nome"]}: {item["vendas"]} unidades', style='List Bullet')
    document.add_heading('Análise de Clientes e Fornecedores', level=1)
    document.add_paragraph(f'• Principal Fornecedor: {data["principalFornecedor"]["nome"]}, Valor Gasto: {data["principalFornecedor"]["valorGasto"]}')
    document.add_heading('Principais Clientes', level=2)
    if data['topClientes']:
        for item in data['topClientes']:
            document.add_paragraph(f'• {item}', style='List Bullet')
    document.save(buffer)
    buffer.seek(0)
@csrf_exempt
@require_POST
@login_required
def api_enviar_relatorios_email(request):
    # ... seu código de enviar email ...
    try:
        data = json.loads(request.body.decode('utf-8'))
        email_destino = data.get('email', '').strip()
        if not email_destino:
            return JsonResponse({'success': False, 'message': 'E-mail não fornecido.'}, status=400)
        total_produtos_qnt = Produto.objects.aggregate(total_qnt=Coalesce(Sum('quantidade'), 0))['total_qnt']
        valor_total_estoque = Produto.objects.aggregate(total_valor=Coalesce(Sum(F('quantidade') * F('preco_venda')), Decimal(0), output_field=DecimalField()))['total_valor']
        produtos_falta = list(Produto.objects.filter(quantidade__lt=F('quantidade_minima')).values_list('nome', flat=True))
        vendas_por_produto = Venda.objects.values('produto__nome').annotate(total_vendido=Coalesce(Sum('quantidade'), 0))
        produtos_mais_saem = [{"nome": item['produto__nome'], "vendas": item['total_vendido']} for item in vendas_por_produto.order_by('-total_vendido')[:3]]
        produtos_menos_saem = [{"nome": item['produto__nome'], "vendas": item['total_vendido']} for item in vendas_por_produto.order_by('total_vendido')[:3]]
        principal_fornecedor = Fornecedor.objects.annotate(valor_gasto=Coalesce(Sum(F('produto__quantidade') * F('produto__preco_compra')), Decimal(0), output_field=DecimalField())).order_by('-valor_gasto').first()
        principal_fornecedor_data = {"nome": principal_fornecedor.nome if principal_fornecedor else "N/A", "valorGasto": f"R$ {principal_fornecedor.valor_gasto:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.') if principal_fornecedor else "R$ 0,00"}
        top_clientes = list(Cliente.objects.annotate(valor_total_vendas=Coalesce(Sum('venda__preco_total'), Decimal(0))).order_by('-valor_total_vendas').values_list('nome', flat=True)[:3])
        fornecedores_cadastrados = Fornecedor.objects.count()
        clientes_cadastrados = Cliente.objects.count()
        total_vendas = Venda.objects.aggregate(total_vendas=Coalesce(Sum('quantidade'), 0))['total_vendas']
        valor_ganho = Venda.objects.aggregate(total_ganho=Coalesce(Sum('preco_total'), Decimal(0), output_field=DecimalField()))['total_ganho']
        dados_relatorio = {
            'produtosEstoque': total_produtos_qnt,
            'valorEstoque': f"R$ {valor_total_estoque:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),
            'produtosFalta': produtos_falta,
            'produtosMaisSaem': produtos_mais_saem,
            'produtosMenosSaem': produtos_menos_saem,
            'principalFornecedor': principal_fornecedor_data,
            'topClientes': top_clientes,
            'fornecedores': fornecedores_cadastrados,
            'clientes': clientes_cadastrados,
            'totalVendas': total_vendas,
            'valorGanho': f"R$ {valor_ganho:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
        }
        buffer_pdf = io.BytesIO()
        generate_pdf(buffer_pdf, dados_relatorio)
        buffer_excel = io.BytesIO()
        generate_excel(buffer_excel, dados_relatorio)
        buffer_word = io.BytesIO()
        generate_word(buffer_word, dados_relatorio)
        subject = "Relatórios Gerados do Sistema"
        message_body = "Olá,\n\nEm anexo, você encontrará os relatórios do sistema em diferentes formatos.\n\nAtenciosamente,\nSua Equipe"
        email = EmailMessage(
            subject,
            message_body,
            settings.EMAIL_HOST_USER,
            [email_destino],
            fail_silently=False,
        )
        email.attach('relatorio_completo.pdf', buffer_pdf.getvalue(), 'application/pdf')
        email.attach('relatorio_completo.xlsx', buffer_excel.getvalue(), 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        email.attach('relatorio_completo.docx', buffer_word.getvalue(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        email.send()
        return JsonResponse({'success': True, 'message': 'Relatórios enviados por e-mail com sucesso!'})
    except json.JSONDecodeError:
        return HttpResponseBadRequest('JSON inválido.')
    except Exception as e:
        print(f"Erro ao enviar relatórios por e-mail: {e}")
        return JsonResponse({'success': False, 'message': f'Erro ao enviar e-mail: {str(e)}'}, status=500)
@login_required
def api_dashboard_data(request):
    try:
        # --- Cálculo dos KPIs dos Cards ---
        total_vendas_valor = Venda.objects.aggregate(
            total=Coalesce(Sum('preco_total'), Decimal(0))
        )['total']
        
        total_vendas_numero = Venda.objects.aggregate(
            total=Coalesce(Count('id'), 0)
        )['total']
        
        total_produtos_estoque = Produto.objects.aggregate(
            total=Coalesce(Sum('quantidade'), 0)
        )['total']
        
        entradas_produtos_mes = HistoricoNotaFiscal.objects.filter(
            data_entrada__year=timezone.now().year,
            data_entrada__month=timezone.now().month
        ).aggregate(
            total=Coalesce(Sum('quantidade_adicionada'), 0)
        )['total']
        
        # --- Dados para Gráficos ---
        hoje = timezone.now().date()
        labels_meses = []
        vendas_mensais = []
        estoque_mensal = []
        sales_growth_labels = []
        sales_growth_data = []
        # Gráfico principal (Vendas e Estoque por mês) e Crescimento de Vendas (Line Chart)
        for i in range(12):
            mes = hoje - timedelta(days=30 * (11 - i))
            
            # Use `calendar` para obter o nome do mês em português
            nome_mes = calendar.month_name[mes.month]
            labels_meses.append(nome_mes[:3].capitalize())
            sales_growth_labels.append(nome_mes[:3].capitalize())
            
            # Vendas (Saídas de Estoque)
            vendas_no_mes = Venda.objects.filter(
                data_venda__year=mes.year,
                data_venda__month=mes.month
            ).aggregate(
                total=Coalesce(Sum('quantidade'), 0)
            )['total']
            vendas_mensais.append(vendas_no_mes)
            
            # Entradas de Estoque
            entradas_no_mes = HistoricoNotaFiscal.objects.filter(
                data_entrada__year=mes.year,
                data_entrada__month=mes.month
            ).aggregate(
                total=Coalesce(Sum('quantidade_adicionada'), 0)
            )['total']
            estoque_mensal.append(entradas_no_mes)
            # Crescimento de Vendas (Line Chart)
            vendas_valor_no_mes = Venda.objects.filter(
                data_venda__year=mes.year,
                data_venda__month=mes.month
            ).aggregate(
                total_venda=Coalesce(Sum('preco_total'), Decimal(0))
            )['total_venda']
            sales_growth_data.append(float(vendas_valor_no_mes))
            
        # Vendas por Produto (Top 5)
        vendas_por_produto_query = Venda.objects.values('produto__nome').annotate(
            total_vendido=Coalesce(Sum('quantidade'), 0)
        ).order_by('-total_vendido')[:5]
        
        vendas_por_produto = {
            'labels': [item['produto__nome'] for item in vendas_por_produto_query],
            'data': [item['total_vendido'] for item in vendas_por_produto_query]
        }
        
        # Compras por Fornecedor (Top 4)
        compras_por_fornecedor_query = Fornecedor.objects.annotate(
            total_gasto=Coalesce(Sum(F('produto__historico_notas__quantidade_adicionada') * F('produto__preco_compra')), Decimal(0))
        ).order_by('-total_gasto')[:4]
        
        compras_por_fornecedor = {
            'labels': [item.nome for item in compras_por_fornecedor_query],
            'data': [float(item.total_gasto) for item in compras_por_fornecedor_query]
        }
        
        # Saída de Produtos (Novo card)
        total_saida_produtos = Venda.objects.aggregate(
            total=Coalesce(Sum('quantidade'), 0)
        )['total']
        # Contas a Pagar
        total_a_pagar = HistoricoNotaFiscal.objects.aggregate(
            total=Coalesce(Sum(F('quantidade_adicionada') * F('preco_compra_nota')), Decimal(0), output_field=DecimalField())
        )['total']
        valor_pago = total_a_pagar * Decimal('0.75')
        total_restante = total_a_pagar - valor_pago
        
        # Criação do dicionário de dados
        dados_dashboard = {
            'kpis': {
                'totalVendas': locale.format_string("%.2f", total_vendas_valor, grouping=True),
                'numeroVendas': total_vendas_numero,
                'produtosEstoque': total_produtos_estoque,
                'entradaProdutos': entradas_produtos_mes
            },
            'gaugeLabels': {
                'saidaProdutos': f"{total_saida_produtos} unidades",
                'pagamentos': f"R$ {total_a_pagar:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
            },
            'mainChart': {
                'labels': labels_meses,
                'vendas': vendas_mensais,
                'estoque': estoque_mensal,
            },
            'inventoryValueChart': {
                'labels': vendas_por_produto['labels'],
                'data': vendas_por_produto['data']
            },
            'daysInventoryChart': {
                'ocupado': total_saida_produtos,
                'livre': 0
            },
            'daysPayableChart': {
                'pago': float(valor_pago),
                'restante': float(total_restante)
            },
            'salesTrendChart': {
                'labels': sales_growth_labels,
                'data': sales_growth_data
            },
            'payableAgingChart': {
                'labels': compras_por_fornecedor['labels'],
                'data': compras_por_fornecedor['data']
            }
        }
        
        return JsonResponse(dados_dashboard)
    except Exception as e:
        print(f"Erro na API de dashboard: {e}")
        return JsonResponse({'error': 'Erro ao carregar dados do dashboard.'}, status=500)
@login_required
def api_keep_alive(request):
    return JsonResponse({'success': True, 'message': 'Sessão estendida.'})
@login_required
def api_historico_notas(request):
    if request.method == 'GET':
        produto_id = request.GET.get('produto_id')
        if not produto_id:
            return JsonResponse({'error': 'ID do produto é obrigatório.'}, status=400)
        try:
            produto = Produto.objects.get(id=produto_id)
            notas = list(HistoricoNotaFiscal.objects.filter(produto=produto).order_by('-data_entrada').values(
                'id', 'numero_nota', 'quantidade_adicionada', 'preco_compra_nota', 'preco_venda_nota', 'data_entrada'
            ))
            return JsonResponse(notas, safe=False)
        except Produto.DoesNotExist:
            return JsonResponse({'error': 'Produto não encontrado.'}, status=404)
    elif request.method == 'POST':
        try:
            data = json.loads(request.body.decode('utf-8'))
            produto_id = data.get('produto_id')
            quantidade_adicionada = int(data.get('quantidade_adicionada'))
            numero_nota = data.get('numero_nota')
            preco_compra_nota = data.get('preco_compra_nota')
            preco_venda_nota = data.get('preco_venda_nota') 
            if not all([produto_id, quantidade_adicionada > 0, numero_nota, preco_compra_nota, preco_venda_nota]):
                return JsonResponse({'error': 'Dados incompletos para criar a nota fiscal.'}, status=400)
            with transaction.atomic():
                produto = get_object_or_404(Produto, id=produto_id)
                produto.quantidade += quantidade_adicionada
                produto.save()
                
                nota_fiscal = HistoricoNotaFiscal.objects.create(
                    produto=produto,
                    numero_nota=numero_nota,
                    quantidade_adicionada=quantidade_adicionada,
                    preco_compra_nota=preco_compra_nota,
                    preco_venda_nota=preco_venda_nota
                )
            return JsonResponse({'message': 'Nota fiscal adicionada com sucesso!', 'nota_id': nota_fiscal.id}, status=201)
        
        except Exception as e:
            return JsonResponse({'error': f'Erro ao adicionar nota: {str(e)}'}, status=400)
            
@login_required
def api_historico_nota_detalhe(request, id):
    try:
        nota = HistoricoNotaFiscal.objects.get(id=id)
    except HistoricoNotaFiscal.DoesNotExist:
        return JsonResponse({'error': 'Nota fiscal não encontrada.'}, status=404)
    if request.method == 'DELETE':
        try:
            with transaction.atomic():
                produto = nota.produto
                if produto.quantidade < nota.quantidade_adicionada:
                    return JsonResponse({'error': f'Não é possível excluir a nota. A quantidade em estoque do produto ficaria negativa. (Removeria: {nota.quantidade_adicionada}, Estoque atual: {produto.quantidade})'}, status=400)
                produto.quantidade -= nota.quantidade_adicionada
                produto.save()
                nota.delete()
            return HttpResponse(status=204)
        except Exception as e:
            return JsonResponse({'error': f'Erro ao excluir nota: {str(e)}'}, status=400)
    return JsonResponse({'error': 'Método não permitido.'}, status=405)

@login_required
def api_cidade_rankings(request):
    """
    API RÁPIDA: Retorna apenas os dados de ranking de cidades.
    """
    try:
        top_clientes_cidades = list(Cliente.objects.exclude(cidade__isnull=True).exclude(cidade__exact='')
                                     .values('cidade', 'estado')
                                     .annotate(count=Count('cpf'))
                                     .order_by('-count')[:5])

        top_fornecedores_cidades = list(Fornecedor.objects.exclude(cidade__isnull=True).exclude(cidade__exact='')
                                         .values('cidade', 'estado')
                                         .annotate(count=Count('cnpj'))
                                         .order_by('-count')[:5])
                                         
        return JsonResponse({
            'top_clientes': top_clientes_cidades,
            'top_fornecedores': top_fornecedores_cidades
        })
    except Exception as e:
        print(f"ERRO ao gerar ranking de cidades: {e}")
        return JsonResponse({'error': 'Erro ao gerar rankings.'}, status=500)


@login_required
def api_map_locations(request):
    """
    API LENTA: Retorna apenas os dados de geolocalização para o mapa usando a API Photon.
    """
    locations = []
    headers = { 'User-Agent': 'SistemaGestao/1.0' }

    # Processa Fornecedores
    for fornecedor in Fornecedor.objects.all():
        if fornecedor.cep and fornecedor.cidade and fornecedor.estado:
            query = f"{fornecedor.cidade}, {fornecedor.estado}, {fornecedor.cep}"
            params = {'q': query, 'limit': 1}
            try:
                time.sleep(1)
                response = requests.get("https://photon.komoot.io/api/", params=params, headers=headers, timeout=10)
                response.raise_for_status()
                data = response.json()
                if data and 'features' in data and data['features']:
                    feature = data['features'][0]
                    lon = feature['geometry']['coordinates'][0]
                    lat = feature['geometry']['coordinates'][1]
                    display_name = f"{feature['properties'].get('name', '')}, {feature['properties'].get('city', fornecedor.cidade)}, {feature['properties'].get('state', fornecedor.estado)}"
                    locations.append({
                        'nome': fornecedor.nome, 'tipo': 'fornecedor', 'latitude': lat,
                        'longitude': lon, 'endereco': display_name.strip(", ")
                    })
            except requests.exceptions.RequestException as e:
                print(f"ERRO DE API (MAPA) para Fornecedor '{fornecedor.nome}': {e}")
    
    # Processa Clientes
    for cliente in Cliente.objects.all():
        if cliente.cep and cliente.cidade and cliente.estado:
            query = f"{cliente.cidade}, {cliente.estado}, {cliente.cep}"
            params = {'q': query, 'limit': 1}
            try:
                time.sleep(1)
                response = requests.get("https://photon.komoot.io/api/", params=params, headers=headers, timeout=10)
                response.raise_for_status()
                data = response.json()
                if data and 'features' in data and data['features']:
                    feature = data['features'][0]
                    lon = feature['geometry']['coordinates'][0]
                    lat = feature['geometry']['coordinates'][1]
                    display_name = f"{feature['properties'].get('name', '')}, {feature['properties'].get('city', cliente.cidade)}, {feature['properties'].get('state', cliente.estado)}"
                    locations.append({
                        'nome': cliente.nome, 'tipo': 'cliente', 'latitude': lat,
                        'longitude': lon, 'endereco': display_name.strip(", ")
                    })
            except requests.exceptions.RequestException as e:
                print(f"ERRO DE API (MAPA) para Cliente '{cliente.nome}': {e}")

    return JsonResponse({'locations': locations})


@login_required
def api_buscar_cep(request, cep):
    """
    Busca um endereço completo a partir de um CEP usando a API ViaCEP.
    """
    try:
        url = f"https://viacep.com.br/ws/{cep}/json/"
        response = requests.get(url, timeout=5)
        response.raise_for_status()
        data = response.json()
        
        if 'erro' not in data:
            return JsonResponse({
                'success': True,
                'endereco': data.get('logradouro', ''),
                'bairro': data.get('bairro', ''),
                'cidade': data.get('localidade', ''),
                'estado': data.get('uf', '')
            })
        else:
            return JsonResponse({'success': False, 'message': 'CEP não encontrado.'}, status=404)
    except (requests.exceptions.RequestException, requests.exceptions.Timeout) as e:
        return JsonResponse({'success': False, 'message': f'Erro na comunicação com a API ViaCEP: {e}'}, status=500)
    except Exception as e:
        return JsonResponse({'success': False, 'message': f'Erro inesperado: {e}'}, status=500)